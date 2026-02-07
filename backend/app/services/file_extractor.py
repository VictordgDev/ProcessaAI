import re
from io import BytesIO

import fitz  # PyMuPDF
from docx import Document


class FileExtractor:
    @staticmethod
    def extract_pdf(content: bytes) -> str:
        """Extrai texto de PDF preservando quebras e marcadores de página."""
        text_parts: list[str] = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text()
                text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                text_parts.append(text)
        return "\n".join(text_parts)

    @staticmethod
    def extract_docx(content: bytes) -> str:
        """Extrai texto de DOCX preservando parágrafos e tabelas."""
        doc = Document(BytesIO(content))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    @staticmethod
    def clean_text(text: str) -> str:
        """Normaliza texto para comparação (sem destruir estrutura)."""
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
